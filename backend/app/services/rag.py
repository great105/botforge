"""
RAG service for BotForge knowledge bases.
Handles: file text extraction, chunking, embedding (OpenRouter), hybrid search.
Adapted from TenderHub's implementation, simplified for bot use case.
"""

import hashlib
import io
import logging
import re
import uuid
from typing import Optional

import httpx
from sqlalchemy import text, select, delete
from sqlalchemy.ext.asyncio import AsyncSession

logger = logging.getLogger(__name__)

OPENROUTER_BASE = "https://openrouter.ai/api/v1"
DEFAULT_EMBEDDING_MODEL = "openai/text-embedding-3-small"
EMBEDDING_DIMS = 1536
CHUNK_TARGET = 350  # tokens
CHUNK_OVERLAP = 70
CHUNK_MIN = 100
CHUNK_MAX = 800


# ─── Text extraction ───────────────────────────────────────

def extract_text_from_bytes(data: bytes, filename: str) -> str:
    """Extract text from file bytes based on extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        return _extract_pdf(data)
    elif ext in ("docx", "doc"):
        return _extract_docx(data)
    elif ext in ("xlsx", "xls"):
        return _extract_xlsx(data)
    elif ext in ("txt", "md", "csv", "tsv", "xml", "json", "html"):
        return _extract_text(data)
    else:
        return _extract_text(data)


def _extract_pdf(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text as pdf_extract
        return pdf_extract(io.BytesIO(data))[:100_000]
    except Exception as e:
        logger.warning(f"PDF extraction failed: {e}")
        return ""


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)[:100_000]
    except Exception as e:
        logger.warning(f"DOCX extraction failed: {e}")
        return ""


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(data), read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)[:100_000]
    except Exception as e:
        logger.warning(f"XLSX extraction failed: {e}")
        return ""


def _extract_text(data: bytes) -> str:
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            return data.decode(enc)[:100_000]
        except (UnicodeDecodeError, ValueError):
            continue
    return ""


# ─── Chunking ──────────────────────────────────────────────

def chunk_text(text_content: str, file_id: str) -> list[dict]:
    """Split text into overlapping chunks. Returns list of {chunk_id, text, index}."""
    if not text_content.strip():
        return []

    paragraphs = re.split(r"\n\s*\n", text_content)
    paragraphs = [p.strip() for p in paragraphs if p.strip()]

    chunks = []
    current = []
    current_len = 0

    for para in paragraphs:
        para_tokens = len(para) // 4  # approximate: 1 token ≈ 4 chars for Russian

        if current_len + para_tokens > CHUNK_MAX and current:
            chunk_text_str = "\n\n".join(current)
            chunks.append(chunk_text_str)
            # Overlap: keep last portion
            overlap_tokens = 0
            overlap_start = len(current)
            for j in range(len(current) - 1, -1, -1):
                overlap_tokens += len(current[j]) // 4
                overlap_start = j
                if overlap_tokens >= CHUNK_OVERLAP:
                    break
            current = current[overlap_start:]
            current_len = sum(len(p) // 4 for p in current)

        current.append(para)
        current_len += para_tokens

        if current_len >= CHUNK_TARGET:
            chunk_text_str = "\n\n".join(current)
            chunks.append(chunk_text_str)
            overlap_tokens = 0
            overlap_start = len(current)
            for j in range(len(current) - 1, -1, -1):
                overlap_tokens += len(current[j]) // 4
                overlap_start = j
                if overlap_tokens >= CHUNK_OVERLAP:
                    break
            current = current[overlap_start:]
            current_len = sum(len(p) // 4 for p in current)

    if current:
        chunk_text_str = "\n\n".join(current)
        if len(chunk_text_str) // 4 >= CHUNK_MIN or not chunks:
            chunks.append(chunk_text_str)
        elif chunks:
            chunks[-1] += "\n\n" + chunk_text_str

    return [
        {"chunk_id": f"{file_id}_{i}", "text": c, "index": i}
        for i, c in enumerate(chunks)
    ]


# ─── Embeddings (OpenRouter) ──────────────────────────────

async def embed_texts(
    texts: list[str],
    api_key: str,
    model: str = DEFAULT_EMBEDDING_MODEL,
) -> list[list[float]]:
    """Generate embeddings via OpenRouter API. Returns list of vectors."""
    if not texts:
        return []

    results = []
    batch_size = 20

    async with httpx.AsyncClient(timeout=30) as client:
        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            try:
                resp = await client.post(
                    f"{OPENROUTER_BASE}/embeddings",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "HTTP-Referer": "https://xn--80acheb8ajrts.xn--p1ai",
                    },
                    json={"model": model, "input": batch},
                )
                resp.raise_for_status()
                data = resp.json()
                for item in sorted(data["data"], key=lambda x: x["index"]):
                    results.append(item["embedding"])
            except Exception as e:
                logger.error(f"Embedding batch {i} failed: {e}")
                results.extend([None] * len(batch))

    return results


# ─── Database operations ───────────────────────────────────

async def ensure_knowledge_table(db: AsyncSession):
    """Create knowledge chunks table if not exists."""
    await db.execute(text("""
        CREATE TABLE IF NOT EXISTS bot_knowledge_chunks (
            id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
            bot_id UUID NOT NULL REFERENCES bots(id) ON DELETE CASCADE,
            file_id VARCHAR(64) NOT NULL,
            filename VARCHAR(255) NOT NULL,
            chunk_index INTEGER NOT NULL,
            chunk_text TEXT NOT NULL,
            embedding vector(1536),
            search_vector tsvector,
            created_at TIMESTAMPTZ DEFAULT now()
        )
    """))
    await db.execute(text("""
        CREATE INDEX IF NOT EXISTS idx_knowledge_bot_id ON bot_knowledge_chunks(bot_id)
    """))
    await db.execute(text("""
        CREATE INDEX IF NOT EXISTS idx_knowledge_file_id ON bot_knowledge_chunks(bot_id, file_id)
    """))
    # HNSW index for vector search (if not exists)
    await db.execute(text("""
        DO $$ BEGIN
            IF NOT EXISTS (SELECT 1 FROM pg_indexes WHERE indexname = 'idx_knowledge_embedding') THEN
                CREATE INDEX idx_knowledge_embedding ON bot_knowledge_chunks
                    USING hnsw (embedding vector_cosine_ops) WITH (m = 16, ef_construction = 64);
            END IF;
        END $$
    """))
    # GIN index for full-text search
    await db.execute(text("""
        CREATE INDEX IF NOT EXISTS idx_knowledge_fts ON bot_knowledge_chunks USING gin(search_vector)
    """))
    await db.commit()


async def store_file_chunks(
    db: AsyncSession,
    bot_id: str,
    file_id: str,
    filename: str,
    chunks: list[dict],
    embeddings: list[Optional[list[float]]],
):
    """Store chunks + embeddings in the database."""
    # Delete old chunks for this file
    await db.execute(
        text("DELETE FROM bot_knowledge_chunks WHERE bot_id = :bid AND file_id = :fid"),
        {"bid": bot_id, "fid": file_id},
    )

    for chunk, embedding in zip(chunks, embeddings):
        emb_str = f"[{','.join(str(x) for x in embedding)}]" if embedding else None
        await db.execute(
            text("""
                INSERT INTO bot_knowledge_chunks
                    (bot_id, file_id, filename, chunk_index, chunk_text, embedding, search_vector)
                VALUES
                    (:bot_id, :file_id, :filename, :idx, :text,
                     CASE WHEN :emb IS NOT NULL THEN CAST(:emb AS vector) ELSE NULL END,
                     to_tsvector('russian', :text) || to_tsvector('english', :text))
            """),
            {
                "bot_id": bot_id,
                "file_id": file_id,
                "filename": filename,
                "idx": chunk["index"],
                "text": chunk["text"],
                "emb": emb_str,
            },
        )
    await db.commit()


async def delete_file_chunks(db: AsyncSession, bot_id: str, file_id: str):
    """Delete all chunks for a file."""
    await db.execute(
        text("DELETE FROM bot_knowledge_chunks WHERE bot_id = :bid AND file_id = :fid"),
        {"bid": bot_id, "fid": file_id},
    )
    await db.commit()


async def list_files(db: AsyncSession, bot_id: str) -> list[dict]:
    """List knowledge base files with chunk counts."""
    result = await db.execute(
        text("""
            SELECT file_id, filename, COUNT(*) as chunks, MIN(created_at) as created_at
            FROM bot_knowledge_chunks
            WHERE bot_id = :bid
            GROUP BY file_id, filename
            ORDER BY MIN(created_at) DESC
        """),
        {"bid": bot_id},
    )
    return [
        {"file_id": r.file_id, "filename": r.filename, "chunks": r.chunks, "created_at": str(r.created_at)}
        for r in result.fetchall()
    ]


async def search_chunks(
    db: AsyncSession,
    bot_id: str,
    query: str,
    api_key: str,
    top_k: int = 5,
    embedding_model: str = DEFAULT_EMBEDDING_MODEL,
) -> list[dict]:
    """Hybrid search: dense (vector) + sparse (FTS), score fusion."""

    # Dense search (vector similarity)
    dense_results = []
    try:
        embeddings = await embed_texts([query], api_key, embedding_model)
        if embeddings and embeddings[0]:
            qvec = f"[{','.join(str(x) for x in embeddings[0])}]"
            result = await db.execute(
                text("""
                    SELECT file_id, filename, chunk_text,
                           1 - (embedding <=> CAST(:qvec AS vector)) AS score
                    FROM bot_knowledge_chunks
                    WHERE bot_id = :bid AND embedding IS NOT NULL
                    ORDER BY embedding <=> CAST(:qvec AS vector)
                    LIMIT :lim
                """),
                {"bid": bot_id, "qvec": qvec, "lim": top_k * 2},
            )
            dense_results = [
                {"file_id": r.file_id, "filename": r.filename, "text": r.chunk_text, "score": float(r.score)}
                for r in result.fetchall()
            ]
    except Exception as e:
        logger.warning(f"Dense search failed: {e}")

    # Sparse search (FTS)
    sparse_results = []
    try:
        result = await db.execute(
            text("""
                SELECT file_id, filename, chunk_text,
                       ts_rank_cd(search_vector,
                           plainto_tsquery('russian', :q) || plainto_tsquery('english', :q), 32) AS score
                FROM bot_knowledge_chunks
                WHERE bot_id = :bid
                  AND search_vector @@ (plainto_tsquery('russian', :q) || plainto_tsquery('english', :q))
                ORDER BY score DESC
                LIMIT :lim
            """),
            {"bid": bot_id, "q": query, "lim": top_k * 2},
        )
        sparse_results = [
            {"file_id": r.file_id, "filename": r.filename, "text": r.chunk_text, "score": float(r.score)}
            for r in result.fetchall()
        ]
    except Exception as e:
        logger.warning(f"Sparse search failed: {e}")

    # Score fusion (RRF)
    scores: dict[str, dict] = {}
    for i, r in enumerate(dense_results):
        key = r["text"][:100]
        scores[key] = {**r, "fused": 0.6 / (i + 1)}
    for i, r in enumerate(sparse_results):
        key = r["text"][:100]
        if key in scores:
            scores[key]["fused"] += 0.4 / (i + 1)
        else:
            scores[key] = {**r, "fused": 0.4 / (i + 1)}

    fused = sorted(scores.values(), key=lambda x: x["fused"], reverse=True)

    if not fused:
        # Fallback: return first chunks
        result = await db.execute(
            text("SELECT filename, chunk_text FROM bot_knowledge_chunks WHERE bot_id = :bid ORDER BY created_at LIMIT :lim"),
            {"bid": bot_id, "lim": top_k},
        )
        return [{"filename": r.filename, "text": r.chunk_text} for r in result.fetchall()]

    return [{"filename": r["filename"], "text": r["text"]} for r in fused[:top_k]]


# ─── High-level operations ─────────────────────────────────

async def process_file(
    db: AsyncSession,
    bot_id: str,
    filename: str,
    file_data: bytes,
    api_key: str,
    embedding_model: str = DEFAULT_EMBEDDING_MODEL,
) -> dict:
    """Process uploaded file: extract → chunk → embed → store. Returns stats."""
    await ensure_knowledge_table(db)

    file_id = hashlib.sha256(f"{bot_id}:{filename}".encode()).hexdigest()[:16]

    # Extract text
    text_content = extract_text_from_bytes(file_data, filename)
    if not text_content.strip():
        return {"error": "Не удалось извлечь текст из файла", "chunks": 0}

    # Chunk
    chunks = chunk_text(text_content, file_id)
    if not chunks:
        return {"error": "Файл пустой или слишком короткий", "chunks": 0}

    # Embed
    embeddings = await embed_texts(
        [c["text"] for c in chunks], api_key, embedding_model
    )

    # Store
    await store_file_chunks(db, bot_id, file_id, filename, chunks, embeddings)

    return {"file_id": file_id, "filename": filename, "chunks": len(chunks)}
